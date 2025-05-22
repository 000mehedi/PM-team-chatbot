from docx import Document
from docx.shared import Inches

# Create a new Word document
doc = Document()

# Title
doc.add_heading('Project Charter: AI Chatbot for Data Insights and Assistance', 0)

# Sections
doc.add_heading('Current Process', level=1)
doc.add_paragraph("""
The Preventative Maintenance (PM) team manages the planning and upkeep of building service equipment across 700+ buildings. Currently, users rely on manual spreadsheet reviews, FAQs, and historical work orders to extract insights. This process is inefficient, especially for non-technical staff, and lacks scalability. Manual searches for equipment manuals based on manufacturer or serial number are time-consuming and often yield inconsistent results.
""")

doc.add_heading('Solution', level=1)
doc.add_paragraph("""
Develop an AI-powered chatbot that allows internal users to:
- Ask questions using natural language.
- Analyze uploaded Excel/CSV files for key metrics and trends.
- Retrieve preventive maintenance FAQs.
- Fetch equipment manuals based on manufacturer and model.
""")

doc.add_heading('Tool Features', level=1)
doc.add_paragraph("""
• Upload Excel/CSV files for instant analysis\n\
• Ask natural language questions about uploaded data\n\
• Generate graphs and summaries automatically\n\
• Save and resume chat sessions\n\
• Restrict access to users with @calgary.ca emails\n\
• Fetch equipment manuals using SerpAPI based on manufacturer/model\n\
• Store data and chat history securely using Supabase
""")

doc.add_heading('Benefits', level=1)
doc.add_paragraph("""
• Saves analysts and engineers hours each week\n\
• Makes data accessible to all team members\n\
• Promotes data-driven decisions\n\
• Reduces repetitive manual tasks\n\
• Fast access to manuals, historical data, and process knowledge
""")

doc.add_heading('Assumptions/Constraints', level=1)
doc.add_paragraph("""
• Users will upload structured Excel or CSV files\n\
• Manual links depend on data availability on third-party sites (e.g., ManualsLib)\n\
• Authentication and user data must be secured (using Supabase & environment variables)\n\
• OpenAI and SerpAPI usage must be monitored for cost control
""")

doc.add_heading('Proof of Concept (PoC) Deliverables', level=1)
doc.add_paragraph("""
1. Frontend: Streamlit user interface\n\
2. Backend: FastAPI (or Streamlit-native), integrated with Supabase\n\
3. OpenAI integration for chatbot responses\n\
4. SerpAPI integration for manual lookups\n\
5. Data visualizations auto-generated from user-uploaded data\n\
6. Secure authentication for Calgary staff only
""")

doc.add_heading('Scope', level=1)
doc.add_paragraph("""
• Provide a web-based chatbot interface for interacting with Excel data\n\
• Add context-aware answers using existing FAQ database\n\
• Deliver basic and advanced charting tools based on user input\n\
• Include ability to fetch manuals via API\n\
• Store session-based memory for chat continuity
""")

doc.add_heading('Schedule', level=1)
doc.add_paragraph("""
• Prompt and interface design: 1–2 weeks\n\
• Endpoint deployment and testing: 1 week\n\
• Streamlit frontend development: 1 week\n\
• Supabase integration and migration: 1–2 weeks\n\
• Full deployment and hosting: 1 week
""")

doc.add_heading('Resources', level=1)
doc.add_paragraph("""
• OpenAI GPT-3.5 or GPT-4 (via Azure)\n\
• SerpAPI for manual lookup\n\
• Supabase for authentication and data storage\n\
• Streamlit Cloud / Railway / Render for deployment\n\
• GitHub for version control\n\
• Docker (optional for scaling)
""")

# Save the document
output_path = "AI_Chatbot_Project_Charter.docx"  # Saves in current directory

doc.save(output_path)

output_path
